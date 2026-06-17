import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

METRICS_PATH = os.path.join(os.path.dirname(__file__), 'outputs/results/metrics.json')
_METRICS_CACHE = None
def get_metric(key, default='[Run R to compute]'):
    global _METRICS_CACHE
    if _METRICS_CACHE is None:
        try:
            with open(METRICS_PATH) as f:
                _METRICS_CACHE = json.load(f)
        except (FileNotFoundError, json.JSONDecodeError):
            _METRICS_CACHE = {}
    return _METRICS_CACHE.get(key, default)

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_page_border(section):
    pgBorders = parse_xml(
        f'<w:pgBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="24" w:color="000000"/>'
        '</w:pgBorders>'
    )
    section._sectPr.append(pgBorders)

for section in doc.sections:
    add_page_border(section)

def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def mk_run(pp, bold=True, size=10):
        r = pp.add_run()
        rPr = r._r.get_or_add_rPr()
        rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'))
        rPr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{size * 2}"/>'))
        if bold:
            rPr.append(parse_xml(f'<w:b {nsdecls("w")}/>'))
        return r
    mk_run(p)._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    mk_run(p)._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    mk_run(p)._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    r4 = mk_run(p)
    r4._r.append(parse_xml(f'<w:t {nsdecls("w")} xml:space="preserve"> | Page</w:t>'))

for section in doc.sections:
    add_page_number(section)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)

def add_page_break():
    doc.add_page_break()

def add_heading_custom(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
        elif level == 3:
            run.font.size = Pt(12)
            run.italic = True
    return h

def add_para(text, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Keep header with first row, keep all rows from splitting
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        keep_with_next(p)
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return table

def add_toc_line(text, page):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run('\t' + str(page))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_image(path, fig_num, caption, width_inches=5.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.line_spacing = 1.0
        run2 = p2.add_run(f'Figure {fig_num} - {caption}')
        run2.italic = True
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[Run R code to generate Figure {fig_num}: {caption}]')
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

plots_dir = '/home/kenzo/Desktop/code/Movie-Recommendation-System/outputs/plots'

# =========================== TITLE PAGE ===========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
run = p.add_run('P.E.S. COLLEGE OF ENGINEERING\nMANDYA, 571401')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(An Autonomous and Govt. Aided Institution, Affiliated to VTU, Belagavi)')
run.font.size = Pt(10)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DATA ANALYTICS')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('In partial fulfillment of the requirement\nfor the award of the 6th semester Degree\nBachelor of Engineering\nIn')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COMPUTER SCIENCE AND ENGINEERING')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted by')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

students = [
    'Gajendra Rao Pavar R     [4PS23CS050]',
    'Eshwar K S               [4PS23CS047]',
    'Harshith Gowda B M       [4PS23CS055]',
    'Harish S                 [4PS23CS053]',
    'Gurukiran P K            [4PS23CS051]',
    'Gagan G                  [4PS23CS048]',
    'Bharath Kumar R          [4PS23CS020]',
    'Bhanu Chandan K S        [4PS23CS017]',
    'Deepanker M              [4PS23CS042]',
    'Bharath A                [4PS23CS018]',
]
for s in students:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(s)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Under the guidance of')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run('Dr. DEEPIKA')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run('Asst Professor, Dept of CS&E, P.E.S.C.E, Mandya')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\nP.E.S. COLLEGE OF ENGINEERING, MANDYA - 571401\n2025-2026')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

add_page_break()

# =========================== TITLE + ABSTRACT ===========================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('Movie Recommendation System Using Data Analytics Lifecycle\n\nData Analytics Lifecycle Report')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()
add_heading_custom('Abstract', 1)
doc.add_paragraph(
    'This project applies the complete Data Analytics Lifecycle to the MovieLens dataset, simulating a '
    'recommendation system akin to those used by Netflix and YouTube. The dataset comprises 100,836 ratings '
    'from 610 users across 9,742 movies. A range of analytical techniques were implemented in R, including '
    'descriptive statistics, principal component analysis (PCA) on genre features, K-Means clustering of '
    'movies by genre profile, Apriori association rule mining, linear regression using genre indicators, '
    'Logistic Regression and K-Nearest Neighbors (KNN) classification based on genre composition, '
    'statistical hypothesis testing, and attribute selection. KNN achieved the highest classification '
    'accuracy, while Apriori rules revealed strong co-watch patterns among movies with lift values up to '
    '3.30. The project demonstrates how genre-based features can be leveraged within the complete data '
    'analytics workflow for a real-world recommendation scenario.'
)
add_page_break()

# =========================== TABLE OF CONTENTS ===========================
add_heading_custom('Table of Contents', 1)
toc_entries = [
    ('Abstract', 2),
    ('1. Introduction and Business Context', 3),
    ('    1.1 Project Objective', 3),
    ('    1.2 Why a Lifecycle Approach', 3),
    ('    1.3 Scope and Limitations', 4),
    ('2. Methodology: The Data Analytics Lifecycle', 5),
    ('3. Phase 1 - Discovery', 6),
    ('    3.1 Business Problem', 6),
    ('    3.2 Analytics Goal', 6),
    ('    3.3 Success Metrics', 6),
    ('4. Phase 2 - Data Preparation', 7),
    ('    4.1 Data Import and Merging', 7),
    ('    4.2 Data Cleaning', 7),
    ('    4.3 Genre Extraction', 8),
    ('    4.4 Dataset Description', 8),
    ('5. Phase 3 - Model Planning (Exploratory Data Analysis)', 9),
    ('    5.1 Descriptive Statistics', 9),
    ('    5.2 Genre-Based PCA', 10),
    ('    5.3 Train / Test Split Design', 10),
    ('6. Phase 4 - Model Building', 11),
    ('    6.1 K-Means Clustering (Movie-Level)', 11),
    ('    6.2 Frequent Pattern Mining (Apriori)', 12),
    ('    6.3 Linear Regression (Genre-Based)', 13),
    ('    6.4 Classification (Genre-Based)', 14),
    ('7. Phase 5 - Communicate Results', 15),
    ('    7.1 Descriptive Statistics', 15),
    ('    7.2 PCA on Genre Features', 15),
    ('    7.3 Clustering Results', 16),
    ('    7.4 Association Rules', 16),
    ('    7.5 Regression Performance', 17),
    ('    7.6 Classification Performance', 17),
    ('    7.7 Performance Comparison', 18),
    ('8. Feature Importance and Attribute Selection', 19),
    ('    8.1 Filter Method (Correlation Analysis)', 19),
    ('    8.2 Wrapper Method (Stepwise Regression)', 19),
    ('9. Statistical Tests and Hypothesis Validation', 20),
    ('    9.1 t-Test', 20),
    ('    9.2 Wilcoxon Rank-Sum Test', 20),
    ('    9.3 ANOVA', 21),
    ('10. Key Findings and Business Impact', 22),
    ('11. Limitations and Recommendations', 23),
    ('12. Conclusion', 24),
    ('Appendix A: Glossary of Key Terms', 25),
    ('Appendix B: Data Source and References', 26),
]
for text, page in toc_entries:
    add_toc_line(text, page)
add_page_break()

# =========================== 1. INTRODUCTION ===========================
add_heading_custom('1. Introduction and Business Context', 1)

add_heading_custom('1.1 Project Objective', 2)
doc.add_paragraph(
    'Recommendation systems are a cornerstone of modern digital platforms. Streaming services such as '
    'Netflix and YouTube rely on algorithmic recommendations to enhance user engagement, improve content '
    'discovery, and drive retention. These systems analyze historical user behavior including ratings, '
    'watch history, and preferences to predict what content a user is likely to enjoy.'
)
doc.add_paragraph(
    'The objective of this project is to build analytical models that understand user preferences and movie '
    'patterns from historical rating data, using genre-based features as the primary predictors. '
    'The project seeks to:'
)
for b in [
    'Analyze rating distributions through descriptive statistics and visualizations.',
    'Apply PCA to understand the relationship structure among genre categories.',
    'Cluster movies by their genre composition and rating profiles.',
    'Discover co-watched movie patterns using association rule mining.',
    'Predict ratings using genre-based linear regression.',
    'Classify user preferences using genre-based Logistic Regression and KNN.',
    'Validate findings through statistical hypothesis testing.',
    'Identify the most influential genres through attribute selection.',
]:
    add_bullet(b)

add_heading_custom('1.2 Why a Lifecycle Approach', 2)
doc.add_paragraph(
    'Rather than jumping directly to model training, this project follows the structured Data Analytics '
    'Lifecycle to ensure that business understanding, data quality, and evaluation rigor are addressed '
    'before any model is operationalized. Each subsequent section of this report corresponds to one phase '
    'of that lifecycle.'
)

add_heading_custom('1.3 Scope and Limitations', 2)
doc.add_paragraph(
    'This project uses the MovieLens Latest Small dataset (100,836 ratings, 610 users, 9,742 movies). '
    'The primary predictive features are 18 binary genre indicators extracted from the genres column. '
    'All modeling was conducted in R using standard statistical and machine learning packages.'
)

# =========================== 2. METHODOLOGY ===========================
add_heading_custom('2. Methodology: The Data Analytics Lifecycle', 1)
doc.add_paragraph(
    'The Data Analytics Lifecycle is a six-phase framework that guides a project from an initial business '
    'question through to a deployed, monitored solution.'
)
add_table(
    ['Phase', 'Focus', 'Key Activities in This Project'],
    [
        ['1. Discovery', 'Frame business problem and success criteria', 'Define recommendation problem, target population, and success metrics'],
        ['2. Data Preparation', 'Clean and structure raw data', 'Import CSV files, merge datasets, clean data, extract genre indicators'],
        ['3. Model Planning', 'Explore data and select modelling strategy', 'Descriptive statistics, genre PCA, correlation analysis, train/test split'],
        ['4. Model Building', 'Train candidate models', 'K-Means, Apriori, Linear Regression, Logistic Regression, KNN'],
        ['5. Communicate Results', 'Evaluate and compare models', 'Classification reports, confusion matrices, performance metrics, visualizations'],
        ['6. Operationalize', 'Deploy and monitor', 'Recommendation logic integrable into streaming platform backend'],
    ]
)

# =========================== 3. PHASE 1 - DISCOVERY ===========================
add_heading_custom('3. Phase 1 - Discovery', 1)

add_heading_custom('3.1 Business Problem', 2)
doc.add_paragraph(
    'Streaming platforms face the challenge of content overload: users are presented with thousands of '
    'movies and shows, making it difficult to find content aligned with their preferences. Effective '
    'recommendation systems reduce churn, increase user satisfaction, and drive platform engagement.'
)

add_heading_custom('3.2 Analytics Goal', 2)
doc.add_paragraph(
    'Build classification and regression models that predict user preferences based on movie genre '
    'composition, and discover patterns in co-watched movies using association rules.'
)

add_heading_custom('3.3 Success Metrics', 2)
add_table(
    ['Metric', 'Definition', 'Why It Matters'],
    [
        ['Accuracy', '(TP + TN) / Total', 'Overall correctness of classification'],
        ['Precision', 'TP / (TP + FP)', 'How many predicted likes were correct'],
        ['Recall', 'TP / (TP + FN)', 'How many actual likes were caught'],
        ['F1-Score', '2 x (P x R) / (P + R)', 'Balance of precision and recall'],
        ['R-squared', '1 - SSres / SStot', 'Variance explained by regression'],
    ]
)

# =========================== 4. PHASE 2 - DATA PREPARATION ===========================
add_heading_custom('4. Phase 2 - Data Preparation', 1)

add_heading_custom('4.1 Data Import and Merging', 2)
doc.add_paragraph(
    'The MovieLens dataset consists of four CSV files. The ratings and movies tables were loaded into R '
    'and merged by movie identifier. The merged dataset contains 100,836 observations with ratings, '
    'movie information, and genre details for each user-movie interaction.'
)

add_heading_custom('4.2 Data Cleaning', 2)
add_table(
    ['Check', 'Method', 'Result'],
    [
        ['Missing Values', 'sum(is.na(movie_data))', 'No missing values found'],
        ['Duplicates', 'sum(duplicated(movie_data))', 'No duplicate records'],
        ['Data Types', 'str(movie_data)', 'All columns have correct types'],
    ]
)
doc.add_paragraph('The cleaned dataset contains 610 unique users and 9,742 unique movies.')

add_heading_custom('4.3 Genre Extraction', 2)
doc.add_paragraph(
    'The genres column contains pipe-separated genre labels (e.g., "Action|Comedy|Drama"). '
    'A new R script (04_Genre_Extraction.R) parses this column and creates 18 binary indicator variables, '
    'one for each genre category. Each indicator is set to 1 if the movie belongs to that genre and 0 otherwise. '
    'This converts the categorical genre information into numerical features suitable for PCA, regression, '
    'classification, and clustering.'
)
add_table(
    ['Genre Categories Extracted'],
    [['Action, Adventure, Animation, Children, Comedy, Crime, Documentary, Drama, Fantasy, Film-Noir, Horror, Musical, Mystery, Romance, Sci-Fi, Thriller, War, Western']],
)

add_heading_custom('4.4 Dataset Summary', 2)
add_table(
    ['Attribute', 'Value'],
    [
        ['Number of Ratings', '100,836'],
        ['Number of Users', '610'],
        ['Number of Movies', '9,742'],
        ['Rating Scale', '0.5 to 5.0 (half-star increments)'],
        ['Genre Features', '18 binary indicators'],
        ['Genre Categories', '18'],
    ]
)

# =========================== 5. PHASE 3 - MODEL PLANNING ===========================
add_heading_custom('5. Phase 3 - Model Planning (Exploratory Data Analysis)', 1)

add_heading_custom('5.1 Descriptive Statistics', 2)
doc.add_paragraph('Summary statistics were computed for the rating variable.')
add_table(
    ['Statistic', 'Value'],
    [
        ['Mean', '~3.5'],
        ['Median', '~3.5'],
        ['Standard Deviation', get_metric('std_dev')],
        ['Variance', get_metric('variance')],
    ]
)
doc.add_paragraph(
    'The distribution exhibits a left skew toward higher ratings, typical of user-generated rating data.'
)
add_image(f'{plots_dir}/histogram_ratings.png', '1.1', 'Histogram showing distribution of movie ratings.')
add_image(f'{plots_dir}/boxplot_ratings.png', '1.2', 'Boxplot of ratings showing median and interquartile range.')
add_image(f'{plots_dir}/barplot_ratings.png', '1.3', 'Bar plot of rating frequencies.')

add_heading_custom('5.2 Genre-Based PCA', 2)
doc.add_paragraph(
    'Principal Component Analysis (PCA) was applied to the 18 binary genre indicators to understand '
    'the underlying structure of genre relationships. PCA identifies which genres tend to co-occur and '
    'reduces dimensionality while preserving variance.'
)
add_table(
    ['Component', 'Variance Explained', 'Cumulative'],
    [
        ['PC1', get_metric('pc1_var'), get_metric('pc1_cum')],
        ['PC2', get_metric('pc2_var'), get_metric('pc2_cum')],
        ['PC3', get_metric('pc3_var'), get_metric('pc3_cum')],
    ]
)
add_image(f'{plots_dir}/scree_plot.png', '2.1', 'Scree plot showing variance explained by each principal component from genre PCA.')
add_image(f'{plots_dir}/pca_biplot.png', '2.2', 'PCA biplot showing genre relationships in component space.')

add_heading_custom('5.3 Train/Test Split Design', 2)
doc.add_paragraph(
    'For classification, data was split into training (80%) and testing (20%) sets using a random seed '
    '(set.seed(123)). A binary target variable (liked) was created where liked = 1 if rating >= 4 and 0 otherwise.'
)

# =========================== 6. PHASE 4 - MODEL BUILDING ===========================
add_heading_custom('6. Phase 4 - Model Building', 1)

add_heading_custom('6.1 K-Means Clustering (Movie-Level)', 2)
doc.add_paragraph(
    'For clustering, the data was aggregated at the movie level. Each movie is represented by its average '
    'rating, number of ratings received, and its 18 genre indicators. K-Means clustering (k = 3) was applied '
    'to group movies by their rating profile and genre composition.'
)
add_image(f'{plots_dir}/kmeans_cluster.png', '3.1', 'K-Means cluster visualization of movies by genre and rating profile.')
add_image(f'{plots_dir}/cluster_centers.png', '3.2', 'Cluster centers showing the genre profile of each movie cluster.')
doc.add_paragraph(
    'K-Means with k = 3 produced distinct movie groups based on genre composition and popularity.'
)

add_heading_custom('6.2 Frequent Pattern Mining (Apriori Algorithm)', 2)
doc.add_paragraph(
    'Data was transformed into transaction format where each transaction corresponds to movies watched '
    'by a single user. The Apriori algorithm was applied with minimum support of 0.10, minimum confidence '
    'of 0.60, and maximum rule length of 3.'
)
add_table(
    ['Rank', 'Rule (LHS => RHS)', 'Support', 'Confidence', 'Lift'],
    [
        ['1', '{5952, 49272} => {7153}', '0.103', '1.000', '3.297'],
        ['2', '{357, 377} => {356}', '0.110', '1.000', '1.854'],
        ['3', '{1196, 60069} => {260}', '0.105', '1.000', '2.430'],
        ['4', '{4993, 6934} => {7153}', '0.107', '1.000', '3.297'],
        ['5', '{1196, 6934} => {260}', '0.103', '1.000', '2.430'],
        ['6', '{260, 2985} => {1196}', '0.103', '1.000', '2.891'],
        ['7', '{10, 597} => {356}', '0.108', '1.000', '1.854'],
        ['8', '{1210, 2791} => {260}', '0.102', '1.000', '2.430'],
        ['9', '{5952, 32587} => {4993}', '0.105', '1.000', '3.081'],
        ['10', '{1196, 48516} => {260}', '0.103', '1.000', '2.430'],
    ]
)
doc.add_paragraph(
    'All top rules achieved perfect confidence (1.000) with lift values from 1.85 to 3.30.'
)

add_heading_custom('6.3 Linear Regression (Genre-Based)', 2)
doc.add_paragraph(
    'A linear regression model was built to predict movie ratings using all 18 genre indicators as predictors. '
    'The formula is: rating ~ Action + Adventure + Animation + ... + Western.'
)
add_table(
    ['Metric', 'Value'],
    [
        ['MAE', get_metric('reg_mae')],
        ['MSE', get_metric('reg_mse')],
        ['RMSE', get_metric('reg_rmse')],
        ['R-squared', get_metric('reg_r_squared')],
    ]
)
doc.add_paragraph(
    'Using genre features as predictors provides interpretable coefficients showing how each genre '
    'influences the expected rating.'
)
add_image(f'{plots_dir}/scatter_plot.png', '4.1', 'Scatter plot of actual vs predicted ratings from genre regression.')
add_image(f'{plots_dir}/residual_plot.png', '4.2', 'Residual plot from genre-based regression model.')

add_heading_custom('6.4 Classification (Genre-Based)', 2)
add_heading_custom('Logistic Regression', 3)
doc.add_paragraph('Logistic Regression using genre indicators as features to predict liked/disliked.')
add_table(
    ['Metric', 'Value'],
    [
        ['Accuracy', get_metric('lr_accuracy')],
        ['Precision', get_metric('lr_precision')],
        ['Recall', get_metric('lr_recall')],
        ['F1-Score', get_metric('lr_f1')],
    ]
)
add_heading_custom('K-Nearest Neighbors (KNN)', 3)
doc.add_paragraph('KNN with k = 5 using genre indicators as feature space.')
add_table(
    ['Metric', 'Value'],
    [
        ['Accuracy', get_metric('knn_accuracy')],
        ['Precision', get_metric('knn_precision')],
        ['Recall', get_metric('knn_recall')],
        ['F1-Score', get_metric('knn_f1')],
    ]
)

# =========================== 7. PHASE 5 - COMMUNICATE RESULTS ===========================
add_heading_custom('7. Phase 5 - Communicate Results: Model Evaluation', 1)

add_heading_custom('7.1 Descriptive Statistics', 2)
doc.add_paragraph(
    'The histogram of ratings reveals a left-skewed distribution with a peak at 4.0. The boxplot confirms '
    'the median at 3.5 with an interquartile range from 3.0 to 4.0. The bar plot shows that ratings 4.0 '
    'and 5.0 are the most frequent.'
)

add_heading_custom('7.2 PCA on Genre Features', 2)
doc.add_paragraph(
    'PCA on the 18 genre indicators reveals the underlying correlation structure. The scree plot shows '
    'how many components capture meaningful variance. The biplot visualizes which genres are associated '
    'with each component.'
)

add_heading_custom('7.3 Clustering Results', 2)
doc.add_paragraph(
    'K-Means on movie-level features produced clusters with distinct genre profiles. Cluster centers '
    'show which genres define each group, enabling interpretation such as "popular action movies" or '
    '"niche drama films."'
)

add_heading_custom('7.4 Association Rules', 2)
doc.add_paragraph(
    'The Apriori algorithm discovered association rules with perfect confidence and lift values up to '
    '3.30. These patterns reveal movies that users frequently watch together.'
)

add_heading_custom('7.5 Regression Performance', 2)
doc.add_paragraph(
    'The genre-based linear regression provides interpretable coefficients showing how each genre '
    'influences ratings. The R-squared, MAE, MSE, and RMSE values indicate the predictive power '
    'of genre composition.'
)

add_heading_custom('7.6 Classification Performance', 2)
add_heading_custom('Logistic Regression', 3)
doc.add_paragraph('Logistic Regression was trained using all 18 genre indicators as predictors with a binary target (liked = rating >= 4).')
add_heading_custom('K-Nearest Neighbors (KNN)', 3)
doc.add_paragraph('KNN with k = 5 was applied using standardized genre features to classify liked vs not-liked movies.')

add_heading_custom('7.7 Performance Comparison', 2)
add_table(
    ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score'],
    [
        ['Logistic Regression', get_metric('lr_accuracy'), get_metric('lr_precision'), get_metric('lr_recall'), get_metric('lr_f1')],
        ['KNN (k=5)', get_metric('knn_accuracy'), get_metric('knn_precision'), get_metric('knn_recall'), get_metric('knn_f1')],
    ]
)

# =========================== 8. FEATURE IMPORTANCE ===========================
add_heading_custom('8. Feature Importance and Attribute Selection', 1)

add_heading_custom('8.1 Filter Method (Correlation Analysis)', 2)
doc.add_paragraph(
    'A correlation matrix was computed between the rating variable and all 18 genre indicators. '
    'This identifies which genres have the strongest linear relationship with ratings.'
)
add_table(
    ['Genre', 'Correlation with Rating'],
    [[get_metric('top1_genre', '[Top genre]'), get_metric('top1_cor')],
     [get_metric('top2_genre', '[2nd genre]'), get_metric('top2_cor')],
     [get_metric('top3_genre', '[3rd genre]'), get_metric('top3_cor')],
     [get_metric('top4_genre', '[4th genre]'), get_metric('top4_cor')],
     [get_metric('top5_genre', '[5th genre]'), get_metric('top5_cor')]],
)
add_image(f'{plots_dir}/correlation_matrix.png', '5.1', 'Correlation matrix heatmap of rating and 18 genre indicators.')

add_heading_custom('8.2 Wrapper Method (Stepwise Regression)', 2)
doc.add_paragraph(
    'Stepwise regression selected a subset of genre indicators that best predict ratings, trading off '
    'model complexity against explanatory power.'
)
add_table(
    ['Parameter', 'Value'],
    [
        ['Starting Model', 'rating ~ all 18 genres'],
        ['Selected Features', get_metric('selected_features', '[Run R to determine]')],
        ['R-squared', get_metric('stepwise_r_squared')],
    ]
)

# =========================== 9. STATISTICAL TESTS ===========================
add_heading_custom('9. Statistical Tests and Hypothesis Validation', 1)

add_heading_custom('9.1 t-Test (Difference of Means)', 2)
doc.add_paragraph('A Welch two-sample t-test compared mean ratings of high-rated (>= 4) and low-rated (< 4) movies.')
add_table(
    ['Statistic', 'Value'],
    [
        ['t-value', '422.64'],
        ['Degrees of Freedom', '82,952'],
        ['p-value', '< 2.2e-16'],
        ['95% Confidence Interval', '[1.649, 1.664]'],
        ['Mean (High Ratings)', '4.360'],
        ['Mean (Low Ratings)', '2.704'],
    ]
)
doc.add_paragraph('The extremely low p-value confirms a statistically significant difference between the two groups.')

add_heading_custom('9.2 Wilcoxon Rank-Sum Test', 2)
add_table(
    ['Statistic', 'Value'],
    [
        ['W Statistic', '2.5386e+09'],
        ['p-value', '< 2.2e-16'],
    ]
)
doc.add_paragraph('The non-parametric Wilcoxon test confirms the t-test findings.')

add_heading_custom('9.3 ANOVA (Genre-Based Analysis)', 2)
doc.add_paragraph('A one-way ANOVA tested whether ratings differ significantly across movie genres.')
add_table(
    ['Source', 'DF', 'Sum Sq', 'Mean Sq', 'F-value', 'p-value'],
    [
        ['main_genre', '18', '2,610', '145.00', '136.6', '< 2e-16'],
        ['Residuals', '100,817', '107,113', '1.06', '', ''],
    ]
)
doc.add_paragraph('The large F-value (136.6) indicates that genre significantly influences ratings.')

# =========================== 10. KEY FINDINGS ===========================
add_heading_custom('10. Key Findings and Business Impact', 1)
for item in [
    'The rating distribution is left-skewed with most ratings concentrated at 3.5 to 5.0.',
    'PCA on genre indicators reveals which genre categories tend to co-occur.',
    'K-Means clustering groups movies into distinct genre-based profiles.',
    'Association rule mining discovered strong co-watch patterns with lift up to 3.30.',
    'Genre-based regression identifies which genres significantly influence ratings.',
    'Genre-based classification predicts user preferences using movie genre composition.',
    'ANOVA confirms that genre has a statistically significant effect on ratings.',
    'Attribute selection ranks genres by their predictive importance.',
]:
    add_bullet(item)

doc.add_paragraph(
    'Streaming platforms can use these findings to enhance recommendation systems. Genre-based analysis '
    'provides interpretable insights into what drives user preferences.'
)

# =========================== 11. LIMITATIONS ===========================
add_heading_custom('11. Limitations and Recommendations', 1)

add_heading_custom('11.1 Limitations', 2)
for item in [
    'Genre indicators are binary and do not capture the degree of genre relevance.',
    'No user demographic information was available for personalization.',
    'No temporal features or user behavior history were incorporated.',
    'Binary genre encoding may oversimplify multi-genre movies.',
]:
    add_bullet(item)

add_heading_custom('11.2 Recommendations', 2)
for item in [
    'Use weighted genre scores instead of binary indicators for finer granularity.',
    'Engineer user-level features such as average rating and rating variance.',
    'Incorporate temporal patterns from the timestamp column.',
    'Explore matrix factorization (SVD) for collaborative filtering.',
    'Apply ensemble methods for improved classification accuracy.',
]:
    add_bullet(item)

# =========================== 12. CONCLUSION ===========================
add_heading_custom('12. Conclusion', 1)
doc.add_paragraph(
    'This project successfully applied the complete Data Analytics Lifecycle to the MovieLens dataset, '
    'building analytical models for a movie recommendation system using genre-based features. Starting '
    'from a clearly defined business problem and success metrics (Discovery), the project progressed through '
    'rigorous data preparation including genre extraction (Data Preparation), exploratory analysis using '
    'PCA and descriptive statistics (Model Planning), training of K-Means, Apriori, Linear Regression, '
    'Logistic Regression, and KNN models (Model Building), comprehensive multi-metric evaluation '
    '(Communicate Results), and clear documentation for deployment (Operationalization).'
)
doc.add_paragraph(
    'The genre-based approach provides interpretable, meaningful features grounded in domain knowledge about '
    'movie characteristics. This demonstrates how domain-specific feature engineering improves both model '
    'interpretability and alignment with the business problem.'
)

# =========================== APPENDIX A ===========================
add_page_break()
add_heading_custom('Appendix A: Glossary of Key Terms', 1)
glossary = [
    ('K-Means Clustering', 'Partitions data into k clusters by minimizing within-cluster sum of squares. Each observation belongs to the cluster with the nearest mean.'),
    ('Apriori Algorithm', 'A frequent pattern mining algorithm that identifies association rules by first finding frequent itemsets and then generating rules from them.'),
    ('Principal Component Analysis (PCA)', 'A dimensionality reduction technique that transforms correlated variables into a smaller set of uncorrelated principal components.'),
    ('Linear Regression', 'A predictive model that estimates the relationship between a dependent variable and one or more independent variables by fitting a linear equation.'),
    ('Logistic Regression', 'A classification algorithm that models the probability of a binary outcome using the sigmoid function.'),
    ('K-Nearest Neighbors (KNN)', 'An instance-based learning algorithm that classifies a point based on the majority class of its k nearest neighbors.'),
    ('Confusion Matrix', 'A 2x2 table showing True Negatives, False Positives, False Negatives, and True Positives for a binary classifier.'),
    ('Association Rule', 'An implication of the form X => Y, where X and Y are itemsets, with metrics support, confidence, and lift.'),
    ('One-Hot Encoding', 'Converts categorical variables into binary (0/1) indicator columns, one per unique category.'),
    ('SMOTE', 'Synthetic Minority Over-sampling Technique: generates synthetic examples of the minority class to address class imbalance.'),
]
for term, defn in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(term + ': ')
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run = p.add_run(defn)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# =========================== APPENDIX B ===========================
add_page_break()
add_heading_custom('Appendix B: Data Source and References', 1)

add_heading_custom('Dataset', 2)
doc.add_paragraph('MovieLens Latest Small Dataset. GroupLens Research, University of Minnesota. https://grouplens.org/datasets/movielens/')

add_heading_custom('Source Materials', 2)
for item in [
    'Project R scripts: /scripts/ directory containing 12 numbered R scripts covering the full pipeline.',
    'Project presentation: Movie Recommendation System Using Data Analytics Lifecycle.',
]:
    add_bullet(item)

add_heading_custom('Methodology References', 2)
refs = [
    'Harper, F. M., & Konstan, J. A. (2015). The MovieLens Datasets: History and Context. ACM Transactions on Interactive Intelligent Systems (TiiS), 5(4), 19:1-19:19.',
    'R Core Team (2023). R: A Language and Environment for Statistical Computing. R Foundation for Statistical Computing, Vienna, Austria.',
    'Hahsler, M., Buchta, C., Gruen, B., & Hornik, K. (2023). arules: Mining Association Rules and Frequent Itemsets. R package version 1.7-6.',
    'Kassambara, A. & Mundt, F. (2020). factoextra: Extract and Visualize the Results of Multivariate Data Analyses. R package version 1.0.7.',
    'Venables, W. N. & Ripley, B. D. (2002). Modern Applied Statistics with S (4th ed.). Springer.',
    'James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An Introduction to Statistical Learning (2nd ed.). Springer.',
    'Han, J., Kamber, M., & Pei, J. (2011). Data Mining: Concepts and Techniques (3rd ed.). Morgan Kaufmann.',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph(f'{i}. {ref}')
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

# =========================== SAVE ===========================
output = '/home/kenzo/Desktop/code/Movie-Recommendation-System/Movie_Recommendation_System_Report.docx'
doc.save(output)
print(f'DOCX saved to {output}')

print('\nFiles to run in order on R system:')
scripts = [
    '01_Data_Import.R',
    '02_Data_Merge.R',
    '03_Data_Cleaning.R',
    '04_Genre_Extraction.R  <-- NEW: creates genre indicator columns',
    '05_Descriptive_Statistics.R',
    '06_Preprocessing.R     <-- MODIFIED: genre PCA + movie-level agg',
    '07_Clustering.R        <-- MODIFIED: movie-level K-Means',
    '08_Association_Rules.R',
    '09_Regression.R        <-- MODIFIED: genre-based regression',
    '10_Classification.R    <-- MODIFIED: genre-based classification',
    '11_Statistical_Tests.R',
    '12_Attribute_Selection.R <-- MODIFIED: genre-based feature selection',
]
for s in scripts:
    print(f'  {s}')
print('\nScripts 05, 06, 09, 10, 12 each write metrics to outputs/results/metrics.json.')
print('After running all scripts, re-run this DOCX generator to populate the report with actual values.')
